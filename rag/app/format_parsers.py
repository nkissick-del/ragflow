#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import re
import logging
from io import BytesIO
from timeit import default_timer as timer
from functools import reduce
from PIL import Image
from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from docx.text.paragraph import Paragraph
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml
from docx.table import Table as DocxTable
from markdown import markdown
from common.constants import LLMType
from api.db.services.llm_service import LLMBundle
from rag.nlp import concat_img, find_codec
from deepdoc.parser import DocxParser, PdfParser, MarkdownParser, PlainParser, MarkdownElementExtractor
from deepdoc.parser.pdf_parser import VisionParser
from deepdoc.parser.figure_parser import vision_figure_parser_pdf_wrapper
from deepdoc.parser.docling_parser import DoclingParser
from deepdoc.parser.tcadp_parser import TCADPParser
import os


# Helper for Docx
def load_from_xml_v2(baseURI, rels_item_xml):
    """
    Return |_SerializedRelationships| instance loaded with the
    relationships contained in *rels_item_xml*. Returns an empty
    collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ("../NULL", "NULL") or rel_elm.target_ref.startswith("#"):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


# Patching _SerializedRelationships.load_from_xml
# This was done inside Chunk in naive.py, but it makes more sense to do it globally or in the module that uses it.
# However, to be safe and match behavior, we can do it when Docx is used or here at module level.
_SerializedRelationships.load_from_xml = load_from_xml_v2


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None
        res_img = None
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            embed = embed[0]
            try:
                related_part = document.part.related_parts[embed]
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logging.info("Unrecognized image format. Skipping image.")
                continue
            except UnexpectedEndOfFileError:
                logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
                continue
            except InvalidImageStreamError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except UnicodeDecodeError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except Exception as e:
                logging.warning(f"The recognized image stream appears to be corrupted. Skipping image, exception: {e}")
                continue
            try:
                image = Image.open(BytesIO(image_blob)).convert("RGB")
                if res_img is None:
                    res_img = image
                else:
                    res_img = concat_img(res_img, image)
            except Exception as e:
                logging.warning(f"Fail to open or concat images, exception: {e}")
                continue

        return res_img

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith("p"):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(("p", i, p))
                elif block.tag.endswith("tbl"):  # Table
                    blocks.append(("t", i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""

        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == "t":
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""  # Target table not found

        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks) - 1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue

            if block_type != "p":
                continue

            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")

        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]

            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks) - 1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue

                    if block_type != "p":
                        continue

                    if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")

                if not found:  # Break if no parent heading is found
                    break

            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        table_idx = 0

        def flush_last_image():
            nonlocal last_image, lines
            if last_image is not None:
                lines.append({"text": "", "image": last_image, "table": None, "style": "Image"})
                last_image = None

        for block in self.doc._element.body:
            if pn > to_page:
                break

            if block.tag.endswith("p"):
                p = Paragraph(block, self.doc)

                if from_page <= pn < to_page:
                    text = p.text.strip()
                    style_name = p.style.name if p.style else ""

                    if text:
                        if style_name == "Caption":
                            former_image = None

                            if lines and lines[-1].get("image") and lines[-1].get("style") != "Caption":
                                former_image = lines[-1].get("image")
                                lines.pop()

                            elif last_image is not None:
                                former_image = last_image
                                last_image = None

                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": former_image if former_image else None,
                                    "table": None,
                                }
                            )

                        else:
                            flush_last_image()
                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": None,
                                    "table": None,
                                }
                            )

                            current_image = self.get_picture(self.doc, p)
                            if current_image is not None:
                                lines.append(
                                    {
                                        "text": "",
                                        "image": current_image,
                                        "table": None,
                                    }
                                )

                    else:
                        current_image = self.get_picture(self.doc, p)
                        if current_image is not None:
                            last_image = current_image

                for run in p.runs:
                    xml = run._element.xml
                    if "lastRenderedPageBreak" in xml:
                        pn += 1
                        continue
                    if "w:br" in xml and 'type="page"' in xml:
                        pn += 1

            elif block.tag.endswith("tbl"):
                if pn < from_page or pn > to_page:
                    table_idx += 1
                    continue

                flush_last_image()
                tb = DocxTable(block, self.doc)
                title = self.__get_nearest_title(table_idx, filename)
                html = "<table>"
                if title:
                    html += f"<caption>Table Location: {title}</caption>"
                for r in tb.rows:
                    html += "<tr>"
                    col_idx = 0
                    try:
                        while col_idx < len(r.cells):
                            span = 1
                            c = r.cells[col_idx]
                            for j in range(col_idx + 1, len(r.cells)):
                                if c.text == r.cells[j].text:
                                    span += 1
                                    col_idx = j
                                else:
                                    break
                            col_idx += 1
                            html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                    except Exception as e:
                        logging.warning(f"Error parsing table, ignore: {e}")
                    html += "</tr>"
                html += "</table>"
                lines.append({"text": "", "image": None, "table": html})
                table_idx += 1

        flush_last_image()
        new_line = [(line.get("text"), line.get("image"), line.get("table")) for line in lines]

        return new_line


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start

        def safe_callback(progress, msg):
            if callback and callable(callback):
                callback(progress, msg)

        safe_callback(0, "OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, safe_callback)
        safe_callback(0, "OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        safe_callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        safe_callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge(zoomin=zoomin)
        safe_callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            self._naive_vertical_merge()
            self._concat_downward()
            # self._final_reading_order_merge()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def md_to_html(self, sections):
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []

        from bs4 import BeautifulSoup

        html_content = markdown(text)
        soup = BeautifulSoup(html_content, "html.parser")
        return soup

    def get_hyperlink_urls(self, soup):
        if soup:
            return {a.get("href") for a in soup.find_all("a") if a.get("href")}
        return set()

    def extract_image_urls_with_lines(self, text):
        md_img_re = re.compile(r"!\[[^\]]*\]\(([^)\s]+)")
        html_img_re = re.compile(r'src=["\\\']([^"\\\'>\\s]+)', re.IGNORECASE)
        urls = []
        seen = set()
        lines = text.splitlines()
        for idx, line in enumerate(lines):
            for url in md_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))
            for url in html_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))

        # cross-line
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(text, "html.parser")
            newline_offsets = [m.start() for m in re.finditer(r"\n", text)] + [len(text)]
            for img_tag in soup.find_all("img"):
                src = img_tag.get("src")
                if not src:
                    continue

                tag_str = str(img_tag)
                pos = text.find(tag_str)
                if pos == -1:
                    # fallback
                    pos = max(text.find(src), 0)
                line_no = 0
                for i, off in enumerate(newline_offsets):
                    if pos <= off:
                        line_no = i
                        break
                if (src, line_no) not in seen:
                    urls.append({"url": src, "line": line_no})
                    seen.add((src, line_no))
        except Exception as e:
            logging.error("Failed to extract image urls: {}".format(e))
            pass

        return urls

    def load_images_from_urls(self, urls, cache=None):
        import requests
        from pathlib import Path

        cache = cache or {}
        images = []
        for url in urls:
            if url in cache:
                if cache[url]:
                    images.append(cache[url])
                continue
            img_obj = None
            try:
                if url.startswith(("http://", "https://")):
                    response = requests.get(url, stream=True, timeout=30)
                    if response.status_code == 200 and response.headers.get("Content-Type", "").startswith("image/"):
                        img_obj = Image.open(BytesIO(response.content)).convert("RGB")
                else:
                    local_path = Path(url)
                    if local_path.exists():
                        img_obj = Image.open(url).convert("RGB")
                    else:
                        logging.warning(f"Local image file not found: {url}")
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
            cache[url] = img_obj
            if img_obj:
                images.append(img_obj)
        return images, cache

    def __call__(self, filename, binary=None, separate_tables=True, delimiter=None, return_section_images=False):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        remainder, tables = self.extract_tables_and_remainder(f"{txt}\n", separate_tables=separate_tables)
        # To eliminate duplicate tables in chunking result, uncomment code below and set separate_tables to True in line 410.
        # extractor = MarkdownElementExtractor(remainder)
        extractor = MarkdownElementExtractor(txt)
        image_refs = self.extract_image_urls_with_lines(txt)
        element_sections = extractor.extract_elements(delimiter, include_meta=True)

        sections = []
        section_images = []
        image_cache = {}
        for element in element_sections:
            content = element["content"]
            start_line = element["start_line"]
            end_line = element["end_line"]
            urls_in_section = [ref["url"] for ref in image_refs if start_line <= ref["line"] <= end_line]
            imgs = []
            if urls_in_section:
                imgs, image_cache = self.load_images_from_urls(urls_in_section, image_cache)
            combined_image = None
            if imgs:
                combined_image = reduce(concat_img, imgs) if len(imgs) > 1 else imgs[0]
            sections.append((content, ""))
            section_images.append(combined_image)

        tbls = []
        for table in tables:
            tbls.append(((None, markdown(table, extensions=["markdown.extensions.tables"])), ""))
        if return_section_images:
            return sections, tbls, section_images
        return sections, tbls


def by_deepdoc(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    callback = callback
    binary = binary
    pdf_parser = pdf_cls() if pdf_cls else Pdf()
    sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)

    tables = vision_figure_parser_pdf_wrapper(
        tbls=tables,
        sections=sections,
        callback=callback,
        **kwargs,
    )
    return sections, tables, pdf_parser


def by_mineru(
    filename,
    binary=None,
    from_page=0,
    to_page=100000,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    mineru_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    pdf_parser = None
    if tenant_id:
        if not mineru_llm_name:
            try:
                from api.db.services.tenant_llm_service import TenantLLMService

                env_name = TenantLLMService.ensure_mineru_from_env(tenant_id)
                candidates = TenantLLMService.query(tenant_id=tenant_id, llm_factory="MinerU", model_type=LLMType.OCR)
                if candidates:
                    mineru_llm_name = candidates[0].llm_name
                elif env_name:
                    mineru_llm_name = env_name
            except Exception as e:  # best-effort fallback
                logging.warning(f"fallback to env mineru: {e}")

        if mineru_llm_name:
            try:
                ocr_model = LLMBundle(tenant_id=tenant_id, llm_type=LLMType.OCR, llm_name=mineru_llm_name, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    lang=lang,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle MinerU ({mineru_llm_name}): {e}")

    if callback:
        callback(-1, "MinerU not found.")
    return None, None, None


def by_docling(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    pdf_parser = DoclingParser()
    parse_method = kwargs.get("parse_method", "raw")

    if not pdf_parser.check_installation():
        if callback and callable(callback):
            callback(-1, "Docling not found.")
        return None, None, pdf_parser

    sections, tables = pdf_parser.parse_pdf(
        filepath=filename,
        binary=binary,
        callback=callback,
        output_dir=os.environ.get("MINERU_OUTPUT_DIR", ""),
        delete_output=bool(int(os.environ.get("MINERU_DELETE_OUTPUT", 1))),
        parse_method=parse_method,
    )
    return sections, tables, pdf_parser


def by_tcadp(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    tcadp_parser = TCADPParser()

    if not tcadp_parser.check_installation():
        if callback and callable(callback):
            callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
        return None, None, tcadp_parser

    sections, tables = tcadp_parser.parse_pdf(filepath=filename, binary=binary, callback=callback, output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""), file_type="PDF")
    return sections, tables, tcadp_parser


def by_paddleocr(
    filename,
    binary=None,
    from_page=0,
    to_page=100000,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    paddleocr_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    pdf_parser = None
    if tenant_id:
        if not paddleocr_llm_name:
            try:
                from api.db.services.tenant_llm_service import TenantLLMService

                env_name = TenantLLMService.ensure_paddleocr_from_env(tenant_id)
                candidates = TenantLLMService.query(tenant_id=tenant_id, llm_factory="PaddleOCR", model_type=LLMType.OCR)
                if candidates:
                    paddleocr_llm_name = candidates[0].llm_name
                elif env_name:
                    paddleocr_llm_name = env_name
            except Exception as e:  # best-effort fallback
                logging.warning(f"fallback to env paddleocr: {e}")

        if paddleocr_llm_name:
            try:
                ocr_model = LLMBundle(tenant_id=tenant_id, llm_type=LLMType.OCR, llm_name=paddleocr_llm_name, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle PaddleOCR ({paddleocr_llm_name}): {e}")
                if callback and callable(callback):
                    callback(-1, f"PaddleOCR parsing failed: {e}")
                return None, None, None

    if callback and callable(callback):
        callback(-1, "PaddleOCR not found.")
    return None, None, None


def by_plaintext(filename, binary=None, from_page=0, to_page=100000, callback=None, **kwargs):
    layout_recognizer = (kwargs.get("layout_recognizer") or "").strip()
    if (not layout_recognizer) or (layout_recognizer == "Plain Text"):
        pdf_parser = PlainParser()
    else:
        tenant_id = kwargs.get("tenant_id")
        if not tenant_id:
            raise ValueError("tenant_id is required when using vision layout recognizer")
        vision_model = LLMBundle(
            tenant_id,
            LLMType.IMAGE2TEXT,
            llm_name=layout_recognizer,
            lang=kwargs.get("lang", "Chinese"),
        )
        pdf_parser = VisionParser(vision_model=vision_model, **kwargs)

    sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
    return sections, tables, pdf_parser


PARSERS = {
    "deepdoc": by_deepdoc,
    "mineru": by_mineru,
    "docling": by_docling,
    "tcadp": by_tcadp,
    "paddleocr": by_paddleocr,
    "plaintext": by_plaintext,  # default
}
